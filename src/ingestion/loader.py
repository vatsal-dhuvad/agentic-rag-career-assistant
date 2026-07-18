import os
import re
import tempfile
from pathlib import Path
from urllib.parse import urljoin, urlparse

from src.config import DEFAULT_RESUME_PATH, FIRESTORE_PORTFOLIO_URL, PORTFOLIO_URL


def clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text or "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(file_path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return clean_text("\n".join(pages))


def extract_text_from_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = clean_text("\n".join(p.text for p in cell.paragraphs))
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

    return clean_text("\n".join(text_parts))


def extract_text_from_upload(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_path = temp_file.name

    try:
        if suffix == ".pdf":
            return extract_text_from_pdf(temp_path)
        if suffix == ".docx":
            return extract_text_from_docx(temp_path)
        if suffix == ".txt":
            return clean_text(Path(temp_path).read_text(encoding="utf-8", errors="ignore"))
        raise ValueError("Please upload PDF, DOCX, or TXT.")
    finally:
        try:
            os.remove(temp_path)
        except OSError:
            pass


def load_default_resume_text() -> str:
    if not DEFAULT_RESUME_PATH.exists():
        return ""
    return extract_text_from_pdf(str(DEFAULT_RESUME_PATH))


def load_portfolio_text(url: str = PORTFOLIO_URL) -> str:
    os.environ.setdefault(
        "USER_AGENT",
        "AgenticRAGCareerAssistant/1.0 portfolio loader by Vatsal Dhuvad",
    )

    try:
        from langchain_community.document_loaders import WebBaseLoader

        loader = WebBaseLoader(
            web_paths=[url],
            header_template={
                "User-Agent": (
                    "Mozilla/5.0 AgenticRAGCareerAssistant/1.0 "
                    "(portfolio loader for career assistant project)"
                )
            },
        )
        docs = loader.load()
    except Exception:
        docs = []

    page_text = "\n\n".join(doc.page_content for doc in docs)
    page_text = clean_text(page_text)
    if len(page_text) > 100:
        return page_text

    firestore_text = load_portfolio_text_from_firestore()
    if firestore_text:
        return firestore_text

    return load_portfolio_text_from_static_site(url)


def build_portfolio_summary(portfolio_text: str) -> str:
    if not portfolio_text.strip():
        return ""

    found_projects = re.findall(
        r"Project: ([^\n]+)\nDescription: ([^\n]+)\nTechnologies: ([^\n]+)",
        portfolio_text,
    )

    if not found_projects:
        return ""

    project_lines = [
        f"- {title.strip()}: {description.strip()} Technologies: {technologies.strip()}."
        for title, description, technologies in found_projects
    ]

    return clean_text(
        "Portfolio project list from Vatsal Dhuvad's website:\n"
        + "\n".join(project_lines)
    )


def build_portfolio_skills_summary(portfolio_text: str) -> str:
    if not portfolio_text.strip():
        return ""

    skill_names = []
    skills_match = re.search(r"Portfolio skills: ([^\n]+)", portfolio_text)
    if skills_match:
        skill_names.extend(split_comma_values(skills_match.group(1)))

    for technologies in re.findall(r"Technologies: ([^\n]+)", portfolio_text):
        skill_names.extend(split_comma_values(technologies))

    unique_skills = []
    seen = set()
    for skill in skill_names:
        normalized = skill.lower()
        if normalized and normalized not in seen:
            seen.add(normalized)
            unique_skills.append(skill)

    if not unique_skills:
        return ""

    return clean_text(
        "Recruiter-friendly skills from Vatsal Dhuvad's portfolio:\n"
        + ", ".join(unique_skills)
    )


def build_portfolio_first_impression(portfolio_text: str) -> str:
    if not portfolio_text.strip():
        return ""

    profile_line = ""
    profile_match = re.search(r"Portfolio profile: ([^\n]+)", portfolio_text)
    if profile_match:
        profile_line = profile_match.group(1).strip()

    skills_summary = build_portfolio_skills_summary(portfolio_text)
    projects_summary = build_portfolio_summary(portfolio_text)

    parts = ["### Portfolio Overview"]

    if profile_line:
        parts.append(profile_line)

    if skills_summary:
        parts.append("### Recruiter-Friendly Skills")
        parts.append(skills_summary.replace("Recruiter-friendly skills from Vatsal Dhuvad's portfolio:", "").strip())

    if projects_summary:
        parts.append("### Portfolio Projects")
        parts.append(projects_summary.replace("Portfolio project list from Vatsal Dhuvad's website:", "").strip())

    return clean_text("\n\n".join(parts))


def split_comma_values(text: str) -> list[str]:
    return [
        item.strip()
        for item in text.split(",")
        if item.strip()
    ]


def load_portfolio_text_from_firestore() -> str:
    try:
        import requests

        response = requests.get(
            FIRESTORE_PORTFOLIO_URL,
            timeout=30,
            headers={"User-Agent": os.environ["USER_AGENT"]},
        )
        response.raise_for_status()
        data = response.json()
    except Exception:
        return ""

    fields = data.get("fields", {})
    if not fields:
        return ""

    portfolio = {
        key: read_firestore_value(value)
        for key, value in fields.items()
    }

    text_parts = []
    profile = portfolio.get("profile", {})
    if profile:
        text_parts.append(
            "Portfolio profile: "
            + " | ".join(
                str(profile.get(key, ""))
                for key in ["name", "title", "email", "phone", "location", "heroBio", "bio"]
                if profile.get(key)
            )
        )

    skills = portfolio.get("skills", [])
    if skills:
        skill_names = [skill.get("name", "") for skill in skills if skill.get("name")]
        text_parts.append("Portfolio skills: " + ", ".join(skill_names))

    projects = portfolio.get("projects", [])
    for project in projects:
        technologies = project.get("technologies", [])
        if isinstance(technologies, list):
            technologies_text = ", ".join(str(item) for item in technologies)
        else:
            technologies_text = str(technologies)

        text_parts.append(
            "\n".join(
                [
                    f"Project: {project.get('title', '')}",
                    f"Description: {project.get('description', '')}",
                    f"Technologies: {technologies_text}",
                    f"Category: {project.get('category', '')}",
                    f"GitHub: {project.get('github', '')}",
                    f"Demo: {project.get('demo', '')}",
                ]
            )
        )

    certificates = portfolio.get("certificates", [])
    if certificates:
        certificate_names = [
            f"{item.get('name', '')} - {item.get('org', '')}"
            for item in certificates
            if item.get("name")
        ]
        text_parts.append("Portfolio certifications: " + "; ".join(certificate_names))

    education = portfolio.get("education", [])
    for item in education:
        text_parts.append(
            "Portfolio education: "
            + " | ".join(
                str(item.get(key, ""))
                for key in ["degree", "institution", "startYear", "endYear", "grade", "description"]
                if item.get(key)
            )
        )

    social_links = portfolio.get("socialLinks", {})
    if social_links:
        text_parts.append(
            "Portfolio links: "
            + " | ".join(f"{key}: {value}" for key, value in social_links.items() if value)
        )

    return clean_text("\n\n".join(text_parts))


def read_firestore_value(value: dict):
    if "stringValue" in value:
        return value["stringValue"]
    if "integerValue" in value:
        return int(value["integerValue"])
    if "doubleValue" in value:
        return float(value["doubleValue"])
    if "booleanValue" in value:
        return value["booleanValue"]
    if "arrayValue" in value:
        return [
            read_firestore_value(item)
            for item in value.get("arrayValue", {}).get("values", [])
        ]
    if "mapValue" in value:
        return {
            key: read_firestore_value(item)
            for key, item in value.get("mapValue", {}).get("fields", {}).items()
        }
    return None


def load_portfolio_text_from_static_site(url: str) -> str:
    try:
        import requests

        response = requests.get(
            url,
            timeout=20,
            headers={"User-Agent": os.environ["USER_AGENT"]},
        )
        response.raise_for_status()
    except Exception:
        return ""

    html_text = response.text
    text_parts = []

    title_match = re.search(r"<title[^>]*>(.*?)</title>", html_text, re.I | re.S)
    if title_match:
        text_parts.append(clean_html_text(title_match.group(1)))

    for content in re.findall(r'<meta[^>]+content=["\']([^"\']+)["\']', html_text, re.I):
        if content:
            text_parts.append(clean_html_text(content))

    script_sources = re.findall(r'<script[^>]+src=["\']([^"\']+)["\']', html_text, re.I)
    for script_src in script_sources:
        script_text = load_script_text(urljoin(url, script_src))
        if script_text:
            text_parts.append(script_text)

    return clean_text("\n".join(dict.fromkeys(text_parts)))


def clean_html_text(text: str) -> str:
    import html

    text = re.sub(r"<[^>]+>", " ", text or "")
    return clean_text(html.unescape(text))


def load_script_text(script_url: str, visited: set[str] | None = None) -> str:
    visited = visited or set()
    if script_url in visited or len(visited) > 8:
        return ""
    visited.add(script_url)

    try:
        import html
        import requests

        response = requests.get(
            script_url,
            timeout=20,
            headers={"User-Agent": os.environ["USER_AGENT"]},
        )
        response.raise_for_status()
    except Exception:
        return ""

    script_text = response.text
    nested_text_parts = []
    nested_assets = re.findall(r'["\']([^"\']*assets/[^"\']+\.js)["\']', script_text)
    parsed_url = urlparse(script_url)
    site_root = f"{parsed_url.scheme}://{parsed_url.netloc}/"
    for asset_path in nested_assets:
        nested_text = load_script_text(urljoin(site_root, asset_path), visited)
        if nested_text:
            nested_text_parts.append(nested_text)

    useful_lines = []
    field_values = re.findall(
        r"(?:name|title|description|bio|heroBio|email|phone|location|category|github|linkedin|demo):\"([^\"]{3,700})\"",
        script_text,
    )
    array_values = re.findall(r'"([^"]{3,120})"', script_text)
    strings = [(value, "") for value in field_values + array_values]
    important_words = (
        "vatsal",
        "portfolio",
        "project",
        "skill",
        "python",
        "machine",
        "learning",
        "data",
        "science",
        "engineer",
        "github",
        "linkedin",
        "resume",
        "certification",
        "education",
        "house",
        "price",
        "predictor",
        "cnn",
        "classification",
        "forecasting",
        "dashboard",
        "recommendation",
        "tensorflow",
        "xgboost",
        "power bi",
    )

    for double_quoted, single_quoted in strings:
        value = html.unescape(double_quoted or single_quoted)
        value = value.replace("\\n", " ").replace("\\t", " ")
        value = re.sub(r"\\u[0-9a-fA-F]{4}", " ", value)
        value = clean_text(value)

        if not value or len(value) < 5:
            continue
        if not any(word in value.lower() for word in important_words):
            continue
        if re.search(r"[{}<>=]{2,}|function|=>|className|children|props", value):
            continue

        useful_lines.append(value)

    return "\n".join(dict.fromkeys(useful_lines + nested_text_parts))
